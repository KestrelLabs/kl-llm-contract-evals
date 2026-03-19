from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "WRITEUP.md"
OUTPUT = ROOT / "WRITEUP_polished.docx"
LOGO_PATH = ROOT.parents[1] / "brand/logo/png/icon-512.png"

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
PAPER_SUBTITLE = "The kestrel-evals Initial Implementation"
AUTHOR_NAME = "Daymian Tomczyk"
AUTHOR_TITLE = "Founder & Principal Engineer"
ORG_NAME = "Kestrel Labs"


@dataclass
class CodeBlock:
    language: str
    content: str


@dataclass
class TableBlock:
    rows: list[list[str]]


@dataclass
class HeadingBlock:
    level: int
    text: str


@dataclass
class ParagraphBlock:
    text: str


@dataclass
class BulletBlock:
    items: list[str]
    ordered: bool = False


@dataclass
class RuleBlock:
    pass


@dataclass
class BlufBlock:
    blocks: list[object]


def set_rfonts(element, font_name: str) -> None:
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), font_name)


def apply_run_font(run, font_name: str, size: int | None = None, *, bold: bool | None = None, italic: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    set_rfonts(run._element, font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def apply_style_font(style, font_name: str, size: int, *, bold: bool = False) -> None:
    style.font.name = font_name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    set_rfonts(style.element, font_name)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_border(paragraph, color: str = "D9D9D9") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        p_bdr.append(el)
    p_pr.append(p_bdr)


def strip_md_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.strip()


def parse_table(lines: list[str], idx: int):
    rows: list[list[str]] = []
    start = idx
    while idx < len(lines):
        line = lines[idx]
        if "|" not in line or not line.strip():
            break
        rows.append([cell.strip() for cell in line.strip().strip("|").split("|")])
        idx += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        header = rows[0]
        body = rows[2:]
        width = len(header)
        normalized = [header[:width]] + [r[:width] + [""] * max(0, width - len(r)) for r in body]
        return TableBlock(normalized), idx
    return None, start


def parse_blocks(text: str) -> list[object]:
    lines = text.splitlines()
    blocks: list[object] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            blocks.append(RuleBlock())
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 2 and text == "BLUF":
                i += 1
                inner: list[object] = []
                buf: list[str] = []
                while i < len(lines):
                    s = lines[i].strip()
                    if s == "---":
                        if buf:
                            inner.append(ParagraphBlock(" ".join(buf).strip()))
                            buf = []
                        break
                    if s.startswith("- "):
                        if buf:
                            inner.append(ParagraphBlock(" ".join(buf).strip()))
                            buf = []
                        items = []
                        while i < len(lines) and lines[i].strip().startswith("- "):
                            items.append(strip_md_inline(lines[i].strip()[2:]))
                            i += 1
                        inner.append(BulletBlock(items))
                        continue
                    if s:
                        buf.append(strip_md_inline(s))
                    i += 1
                if buf:
                    inner.append(ParagraphBlock(" ".join(buf).strip()))
                blocks.append(BlufBlock(inner))
                continue
            blocks.append(HeadingBlock(level, strip_md_inline(text)))
            i += 1
            continue

        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            i += 1
            content_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                content_lines.append(lines[i].rstrip("\n"))
                i += 1
            if i < len(lines):
                i += 1
            blocks.append(CodeBlock(lang, "\n".join(content_lines).rstrip()))
            continue

        table_block, new_i = parse_table(lines, i)
        if table_block is not None:
            blocks.append(table_block)
            i = new_i
            continue

        if re.match(r"^-\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^-\s+", lines[i].strip()):
                items.append(strip_md_inline(re.sub(r"^-\s+", "", lines[i].strip())))
                i += 1
            blocks.append(BulletBlock(items))
            continue

        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(strip_md_inline(re.sub(r"^\d+\.\s+", "", lines[i].strip())))
                i += 1
            blocks.append(BulletBlock(items, ordered=True))
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(strip_md_inline(lines[i].strip()[1:].strip()))
                i += 1
            blocks.append(ParagraphBlock(" ".join(quote_lines)))
            continue

        para = [strip_md_inline(stripped)]
        i += 1
        while i < len(lines):
            s = lines[i].strip()
            if not s:
                break
            if s == "---" or s.startswith("#") or s.startswith("```") or s.startswith("- ") or re.match(r"^\d+\.\s+", s) or "|" in s:
                break
            para.append(strip_md_inline(s))
            i += 1
        blocks.append(ParagraphBlock(" ".join(para).strip()))
    return blocks


def load_font(size: int, bold: bool = False):
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSerif-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSerif-Regular.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


def create_architecture_figure(output_path: Path) -> None:
    width, height = 1500, 620
    bg = (255, 255, 255)
    ink = (32, 32, 32)
    border = (110, 110, 110)
    accent = (239, 243, 251)

    img = Image.new("RGB", (width, height), bg)
    draw = ImageDraw.Draw(img)
    box_font = load_font(22, bold=True)
    small_font = load_font(18, bold=False)
    title_font = load_font(26, bold=True)

    draw.text((60, 28), "Figure 1. Architecture overview", fill=ink, font=title_font)

    boxes = [
        ((60, 110, 330, 200), "Suite YAML\nexamples/structured_extraction.yaml"),
        ((400, 110, 640, 200), "suite_loader.py\nvalidated EvalSuite"),
        ((710, 110, 910, 200), "runner.py"),
        ((980, 110, 1200, 200), "OpenAIProvider\n(generate)"),
        ((1270, 110, 1435, 200), "LLM output"),
        ((980, 320, 1200, 420), "checks.py\njson_schema\nallowed_values\nregex"),
        ((1270, 320, 1435, 420), "JSON report"),
        ((1270, 485, 1435, 570), "CLI exit code / CI gate"),
    ]

    def center_text(box, text, font):
        x1, y1, x2, y2 = box
        lines = text.split("\n")
        total_h = sum(font.getbbox(line)[3] - font.getbbox(line)[1] for line in lines) + 8 * (len(lines) - 1)
        y = y1 + ((y2 - y1) - total_h) / 2
        for line in lines:
            bbox = font.getbbox(line)
            w = bbox[2] - bbox[0]
            h = bbox[3] - bbox[1]
            x = x1 + ((x2 - x1) - w) / 2
            draw.text((x, y), line, fill=ink, font=font)
            y += h + 8

    for box, text in boxes:
        draw.rounded_rectangle(box, radius=16, outline=border, width=3, fill=accent)
        center_text(box, text, box_font)

    def arrow(start, end):
        draw.line([start, end], fill=border, width=4)
        dx = end[0] - start[0]
        dy = end[1] - start[1]
        if abs(dx) >= abs(dy):
            if dx >= 0:
                pts = [(end[0], end[1]), (end[0] - 14, end[1] - 8), (end[0] - 14, end[1] + 8)]
            else:
                pts = [(end[0], end[1]), (end[0] + 14, end[1] - 8), (end[0] + 14, end[1] + 8)]
        else:
            if dy >= 0:
                pts = [(end[0], end[1]), (end[0] - 8, end[1] - 14), (end[0] + 8, end[1] - 14)]
            else:
                pts = [(end[0], end[1]), (end[0] - 8, end[1] + 14), (end[0] + 8, end[1] + 14)]
        draw.polygon(pts, fill=border)

    arrow((330, 155), (400, 155))
    arrow((640, 155), (710, 155))
    arrow((910, 155), (980, 155))
    arrow((1200, 155), (1270, 155))
    arrow((1090, 200), (1090, 320))
    arrow((1200, 370), (1270, 370))
    arrow((1350, 420), (1350, 485))

    draw.text((1000, 255), "deterministic validation", fill=ink, font=small_font)
    draw.text((1215, 520), "non-zero on failure", fill=ink, font=small_font)

    img.save(output_path)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.different_first_page_header_footer = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh = header.add_run("kestrel-evals")
    apply_run_font(rh, BODY_FONT, 9, italic=True, color=RGBColor(96, 96, 96))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)
    for run in footer.runs:
        apply_run_font(run, BODY_FONT, 9, color=RGBColor(96, 96, 96))

    apply_style_font(doc.styles["Normal"], BODY_FONT, 11)
    apply_style_font(doc.styles["Title"], BODY_FONT, 18, bold=True)
    apply_style_font(doc.styles["Heading 1"], BODY_FONT, 13, bold=True)
    apply_style_font(doc.styles["Heading 2"], BODY_FONT, 12, bold=True)
    apply_style_font(doc.styles["Heading 3"], BODY_FONT, 11, bold=True)


def add_title_page(doc: Document, title: str) -> None:
    if LOGO_PATH.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_before = Pt(42)
        logo_run = logo_p.add_run()
        logo_run.add_picture(str(LOGO_PATH), width=Inches(0.8))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title)
    apply_run_font(r, BODY_FONT, 18, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(PAPER_SUBTITLE)
    apply_run_font(r2, BODY_FONT, 11, italic=True, color=RGBColor(96, 96, 96))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(30)
    r3 = p3.add_run(AUTHOR_NAME)
    apply_run_font(r3, BODY_FONT, 11, bold=True)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(AUTHOR_TITLE)
    apply_run_font(r4, BODY_FONT, 10, color=RGBColor(64, 64, 64))

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run(ORG_NAME)
    apply_run_font(r5, BODY_FONT, 11)

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r6 = p6.add_run("March 2026")
    apply_run_font(r6, BODY_FONT, 10, color=RGBColor(96, 96, 96))

    doc.add_page_break()


def render_bluf(doc: Document, block: BlufBlock) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F7FB")

    p = cell.paragraphs[0]
    r = p.add_run("BLUF")
    apply_run_font(r, BODY_FONT, 11, bold=True)

    for inner in block.blocks:
        if isinstance(inner, ParagraphBlock):
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(inner.text)
            apply_run_font(r, BODY_FONT, 11)
        elif isinstance(inner, BulletBlock):
            for item in inner.items:
                p = cell.add_paragraph(style="List Bullet")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                r = p.add_run(item)
                apply_run_font(r, BODY_FONT, 10)
    doc.add_paragraph()


def render_heading(doc: Document, block: HeadingBlock) -> None:
    if block.level == 1:
        return
    style = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}.get(block.level, "Heading 3")
    p = doc.add_paragraph(block.text, style=style)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def render_paragraph(doc: Document, text: str, *, indent: bool) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.25) if indent else Inches(0)
    p.paragraph_format.line_spacing = 1.15
    for part in re.split(r"(`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            apply_run_font(r, MONO_FONT, 9)
        else:
            r = p.add_run(part)
            apply_run_font(r, BODY_FONT, 11)


def render_bullets(doc: Document, block: BulletBlock) -> None:
    style = "List Number" if block.ordered else "List Bullet"
    for item in block.items:
        p = doc.add_paragraph(style=style)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(item)
        apply_run_font(r, BODY_FONT, 10)


def render_code(doc: Document, block: CodeBlock, tmpdir: Path) -> None:
    if block.language == "mermaid":
        figure_path = tmpdir / "architecture_diagram.png"
        create_architecture_figure(figure_path)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(figure_path), width=Inches(6.4))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        r = cap.add_run("Figure 1. Architecture overview")
        apply_run_font(r, BODY_FONT, 10, italic=True)
        return

    for line in block.content.splitlines() or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.right_indent = Inches(0.2)
        set_paragraph_border(p, color="CFCFCF")
        r = p.add_run(line)
        apply_run_font(r, MONO_FONT, 9, color=RGBColor(44, 44, 44))
    doc.add_paragraph()


def render_table(doc: Document, block: TableBlock) -> None:
    if not block.rows:
        return
    rows = len(block.rows)
    cols = max(len(r) for r in block.rows)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.autofit = True
    for i, row in enumerate(block.rows):
        for j in range(cols):
            text = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            if i == 0:
                set_cell_shading(cell, "EFEFEF")
            cell.text = text
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    apply_run_font(r, BODY_FONT, 10, bold=(i == 0))
    doc.add_paragraph()


def main() -> None:
    text = INPUT.read_text(encoding="utf-8")
    blocks = parse_blocks(text)
    if not blocks or not isinstance(blocks[0], HeadingBlock):
        raise RuntimeError("Expected first block to be H1 title")
    title = blocks[0].text

    doc = Document()
    style_document(doc)
    add_title_page(doc, title)

    with tempfile.TemporaryDirectory() as td:
        tmpdir = Path(td)
        indent_next_paragraph = False
        for block in blocks[1:]:
            if isinstance(block, BlufBlock):
                render_bluf(doc, block)
                indent_next_paragraph = False
            elif isinstance(block, HeadingBlock):
                render_heading(doc, block)
                indent_next_paragraph = False
            elif isinstance(block, ParagraphBlock):
                render_paragraph(doc, block.text, indent=indent_next_paragraph)
                indent_next_paragraph = True
            elif isinstance(block, BulletBlock):
                render_bullets(doc, block)
                indent_next_paragraph = False
            elif isinstance(block, CodeBlock):
                render_code(doc, block, tmpdir)
                indent_next_paragraph = False
            elif isinstance(block, TableBlock):
                render_table(doc, block)
                indent_next_paragraph = False
            elif isinstance(block, RuleBlock):
                doc.add_paragraph()
                indent_next_paragraph = False

        doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
